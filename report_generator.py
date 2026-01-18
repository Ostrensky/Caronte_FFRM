# --- FILE: app/report_generator.py ---

from docxtpl import DocxTemplate
import docx
from docx2pdf import convert
import os
from app.config import get_custom_general_texts, get_custom_auto_texts, DEFAULT_GENERAL_TEXTS
from document_parts import formatar_motivo_detalhado, create_table_for_auto
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt # ✅ Ensure this is imported
import logging 
from jinja2 import Environment, ChainableUndefined
import time
import re # ✅ Added re for year extraction

def _safe_save(doc, path, retries=5, delay=0.5):
    """
    Tenta salvar um ficheiro docx, com repetições em caso de PermissionError.
    """
    for i in range(retries):
        try:
            doc.save(path)
            logging.info(f"SaveWrapper: Ficheiro salvo com sucesso na tentativa {i+1}: {path}")
            return 
        except PermissionError as e:
            logging.warning(f"SaveWrapper: Tentativa {i+1} falhou. O ficheiro está bloqueado: {path}. A tentar novamente em {delay}s...")
            time.sleep(delay) 
        except Exception as e:
            logging.error(f"SaveWrapper: Falha ao salvar com erro não-relacionado à permissão: {e}")
            raise e 
    
    logging.error(f"SaveWrapper: Falha ao salvar após {retries} tentativas: {path}. O ficheiro continua bloqueado.")
    raise PermissionError(f"Não foi possível salvar o ficheiro, ele continua bloqueado: {path}")

def _format_currency(value):
    """Helper function to format a number into BRL currency string."""
    if isinstance(value, (int, float)):
        return f"R$ {value:,.2f}".replace(",", "v").replace(".", ",").replace("v", ".")
    return str(value)

def _delete_paragraph(paragraph):
    """Helper to delete a paragraph from a python-docx document."""
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

def create_dams_table(doc, pagamentos_list):
    """
    Creates the table for 'Pagamentos Avulsos' (DAMs).
    """
    if not pagamentos_list:
        return None

    # Define headers (6 Columns)
    headers = ['Código Verificação', 'Competência', 'Receita', 'Valor Pago', 'Tributo', 'Notas Associadas']
    
    table = doc.add_table(rows=1, cols=len(headers), style='Table Grid')
    table.autofit = True

    # Setup Header Row
    hdr_cells = table.rows[0].cells
    for i, header_text in enumerate(headers):
        hdr_cells[i].text = header_text
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Populate Data
    for pag in pagamentos_list:
        row_cells = table.add_row().cells
        # Use .get() to avoid errors if key is missing
        row_cells[0].text = str(pag.get('codigo', ''))
        row_cells[1].text = str(pag.get('competencia', ''))
        row_cells[2].text = str(pag.get('receita', '')) 
        row_cells[3].text = str(pag.get('valor_pago', ''))
        row_cells[4].text = str(pag.get('tributo', ''))
        row_cells[5].text = str(pag.get('notas_associadas', ''))

        # Center align everything
        for cell in row_cells:
            cell.vertical_alignment = 1 # Center
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    return table

def create_conclusion_table(doc, summary_data, idd_mode=False):
    if not summary_data:
        return None

    table = doc.add_table(rows=1, cols=7, style='Table Grid')
    table.autofit = True

    # Headers
    hdr_cells = table.rows[0].cells
    headers = ['Tipo', 'Número', 'Exercício', 'NFS-e tributadas', 'ISS - Valor Original', 'Total Crédito Tributário', 'Motivo']
    
    for i, header_text in enumerate(headers):
        hdr_cells[i].text = header_text
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for auto in summary_data.get('autos', []):
        row_cells = table.add_row().cells
        
        # Determine Label
        motive_text = str(auto.get('motivo', '')).strip()
        motive_upper = motive_text.upper()
        
        if "IDD" in motive_upper:
            label_tipo = "IDD"
        elif idd_mode:
            label_tipo = "IDD"
        else:
            label_tipo = "Auto de Infração"
            
        year_text = "-"
        match = re.search(r'\((\d{4})\)', motive_text)
        if match:
            year_text = match.group(1)
        
        row_cells[0].text = label_tipo
        row_cells[1].text = str(auto.get('numero', ''))
        row_cells[2].text = year_text 
        row_cells[3].text = str(auto.get('nfs_tributadas', ''))
        row_cells[4].text = _format_currency(auto.get('iss_valor_original', 0))
        row_cells[5].text = _format_currency(auto.get('total_credito_tributario', 0)) 
        row_cells[6].text = motive_text

    # ✅ CHANGED: Loop over Multiple Multas
    multas_list = summary_data.get('multas', [])
    if multas_list and not idd_mode:
        for m in multas_list:
            row_cells = table.add_row().cells
            row_cells[0].text = 'Multa'
            row_cells[0].paragraphs[0].runs[0].font.bold = True
            row_cells[1].text = str(m.get('numero', ''))
            row_cells[2].text = str(m.get('ano', '-')) # ✅ Use specific year
            row_cells[3].text = '-'
            row_cells[4].text = '-'
            row_cells[5].text = _format_currency(m.get('valor_credito', 0))
            row_cells[6].text = 'Descumprimento de dever instrumental'

    # --- Total Row ---
    total_row_cells = table.add_row().cells
    total_row_cells[0].text = 'Total'
    total_row_cells[0].paragraphs[0].runs[0].font.bold = True
    total_row_cells[1].text = '-'
    total_row_cells[2].text = '-' 
    total_row_cells[3].text = '-'
    total_row_cells[4].text = '-'
    total_row_cells[5].text = _format_currency(summary_data.get('total_geral_credito', 0))
    total_row_cells[5].paragraphs[0].runs[0].font.bold = True
    total_row_cells[6].text = '-'

    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i != 6:
                 cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    return table

def convert_to_pdf(docx_path):
    """
    Converts a specific DOCX file to PDF using docx2pdf (requires MS Word on Windows).
    """
    try:
        logging.info(f"PDF Convert: Iniciando conversão para {docx_path}...")
        pdf_path = docx_path.replace(".docx", ".pdf")
        
        # Check if output file already exists and is locked
        if os.path.exists(pdf_path):
            try:
                os.remove(pdf_path)
            except PermissionError:
                logging.error(f"PDF Convert: O ficheiro de destino está aberto/bloqueado: {pdf_path}")
                return None

        convert(docx_path, pdf_path)
        logging.info(f"PDF Convert: Sucesso! Salvo em {pdf_path}")
        return pdf_path
    except Exception as e:
        logging.error(f"PDF Convert: Falha ao converter {docx_path}. Erro: {e}")
        return None

def generate_report(context, template_path, output_path, temp_path):
    try:
        custom_general_texts = get_custom_general_texts()
        custom_auto_texts = get_custom_auto_texts()
        jinja_env = Environment(undefined=ChainableUndefined)

        raw_idd_mode = context.get('idd_mode', False)
        idd_mode = str(raw_idd_mode).lower() in ('true', '1', 't') if isinstance(raw_idd_mode, str) else bool(raw_idd_mode)
        logging.info(f"Report Gen: IDD Mode is {idd_mode}")

        rendered_general_texts = {}
        for key, template_string in custom_general_texts.items():
            try:
                template = jinja_env.from_string(template_string)
                rendered_general_texts[key] = template.render(context)
            except Exception:
                rendered_general_texts[key] = f"[[ERRO: {key}]]"

        render_context = rendered_general_texts.copy()
        render_context.update(context) 

        # --- TÍTULO E INTRO ---
        if idd_mode:
            if 'TITULO_DOCUMENTO_IDD' in rendered_general_texts:
                render_context['titulo_documento'] = rendered_general_texts['TITULO_DOCUMENTO_IDD']
            elif 'TITULO_DOCUMENTO_IDD' in DEFAULT_GENERAL_TEXTS:
                render_context['titulo_documento'] = DEFAULT_GENERAL_TEXTS['TITULO_DOCUMENTO_IDD']
            else:
                render_context['titulo_documento'] = "Informação Fiscal - Operação IDD"

            if 'I_INTRO_IDD' in rendered_general_texts:
                render_context['I_INTRO'] = rendered_general_texts['I_INTRO_IDD']
            elif 'I_INTRO_IDD' in DEFAULT_GENERAL_TEXTS:
                try:
                    tmpl = jinja_env.from_string(DEFAULT_GENERAL_TEXTS['I_INTRO_IDD'])
                    render_context['I_INTRO'] = tmpl.render(context)
                except:
                    render_context['I_INTRO'] = DEFAULT_GENERAL_TEXTS['I_INTRO_IDD']

            render_context['V_CONCLUSAO_FINAL'] = "Os créditos confessados foram devidamente formalizados, conforme explicações anteriores."
            render_context['v_conclusao_intro'] = ""

            if 'multa_aplicada' in render_context:
                render_context['multa_aplicada']['aplicada'] = False
                render_context['multa_aplicada']['texto_multa'] = ""
            
            render_context['multa_dispensada_is'] = False
            render_context['multa_dispensada_simples'] = False
            render_context['multa_sem_infracao'] = False

        else:
            if 'TITULO_DOCUMENTO_AUTO' in rendered_general_texts:
                render_context['titulo_documento'] = rendered_general_texts['TITULO_DOCUMENTO_AUTO']
            elif 'TITULO_DOCUMENTO_AUTO' in DEFAULT_GENERAL_TEXTS:
                render_context['titulo_documento'] = DEFAULT_GENERAL_TEXTS['TITULO_DOCUMENTO_AUTO']

            if 'V_CONCLUSAO_INTRO_RECEITAS' in rendered_general_texts:
                render_context['v_conclusao_intro'] = rendered_general_texts['V_CONCLUSAO_INTRO_RECEITAS']
            else:
                render_context['v_conclusao_intro'] = "Em suma, os seguintes lançamentos foram realizados nesta Operação Receita:"

        render_context.pop('autos', None)
        render_context.pop('summary', None)

        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Template file not found: {template_path}")
            
        doc_tpl = DocxTemplate(template_path)
        doc_tpl.tpl_jinja_env = jinja_env
        
        doc_tpl.render(render_context)
        _safe_save(doc_tpl, temp_path)

        doc = docx.Document(temp_path)
        
        term_label = "IDD" if idd_mode else "Auto de Infração"

        if idd_mode:
            paragraphs_to_remove = []
            for p in doc.paragraphs:
                if "IV - DEVERES INSTRUMENTAIS" in p.text or "IV – DEVERES INSTRUMENTAIS" in p.text:
                    paragraphs_to_remove.append(p)
            for p in paragraphs_to_remove:
                _delete_paragraph(p)

            for p in doc.paragraphs:
                text_clean = p.text.strip()
                if "V – CONCLUSÃO" in text_clean or "V - CONCLUSÃO" in text_clean:
                    replaced_in_run = False
                    for run in p.runs:
                        if "V – CONCLUSÃO" in run.text:
                            run.text = run.text.replace("V – CONCLUSÃO", "IV – CONCLUSÃO")
                            replaced_in_run = True
                        elif "V - CONCLUSÃO" in run.text:
                            run.text = run.text.replace("V - CONCLUSÃO", "IV – CONCLUSÃO")
                            replaced_in_run = True
                    if not replaced_in_run:
                        p.text = p.text.replace("V – CONCLUSÃO", "IV – CONCLUSÃO").replace("V - CONCLUSÃO", "IV – CONCLUSÃO")

        # --- INSERÇÃO DINÂMICA DE TABELAS ---

        # 1. Autos Table
        placeholder_p = next((p for p in doc.paragraphs if '###INFRACTIONS_SECTION_PLACEHOLDER###' in p.text), None)
        if placeholder_p:
            parent = placeholder_p._p.getparent()
            autos_list = context.get('autos', [])

            def get_year_from_auto(auto):
                # Extract year from motive string e.g. "Alíquota (2022)"
                # or fallback to 0 to keep them at the top
                motive_str = str(auto.get('motivo', {}).get('texto_simples', '') or auto.get('motivo', ''))
                # Also check top-level keys if 'motivo' is dict
                match = re.search(r'\((\d{4})\)', motive_str)
                if match:
                    return int(match.group(1))
                # Fallback: check raw string if 'motivo' is just a string
                match_raw = re.search(r'\((\d{4})\)', str(auto.get('motive_text', '')))
                if match_raw:
                    return int(match_raw.group(1))
                return 0

            # Sort by Year, then by Auto Number
            autos_list.sort(key=lambda x: (get_year_from_auto(x), x.get('numero', '')))
            
            current_year_header = None


            for auto_data in autos_list:
                # ✅ STEP B: Insert Year Header if Changed
                this_year = get_year_from_auto(auto_data)
                
                if this_year > 0 and this_year != current_year_header:
                    current_year_header = this_year
                    
                    # Create Subtitle Paragraph
                    year_p = doc.add_paragraph()
                    run = year_p.add_run(f"Exercício {this_year}")
                    run.bold = True
                    run.font.size = Pt(12)
                    year_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                    # Insert before placeholder
                    parent.insert(parent.index(placeholder_p._p), year_p._p)

                # ✅ STEP C: Insert Auto Details (Existing Logic)
                motivo_formatado = formatar_motivo_detalhado(auto_data.get('motivo', {}), custom_auto_texts)
                
                # Check for IDD label override in motive text
                # (Reusing logic from summary table if desired, or stick to term_label)
                final_label = term_label
                if "IDD" in str(auto_data.get('motive_text', '')).upper():
                    final_label = "IDD"

                intro_text = (
                    f"·   {final_label} {auto_data.get('numero', 'N/A')} = "
                    f"NFS-e de nº(s) {auto_data.get('nfs_e_numeros', '[N/A]')} – "
                    f"período de {auto_data.get('periodo', '[N/A]')} – {motivo_formatado}"
                )
                
                intro_p = doc.add_paragraph(intro_text)
                table = create_table_for_auto(doc, auto_data, idd_mode=idd_mode)
                spacer_p = doc.add_paragraph()
                
                parent.insert(parent.index(placeholder_p._p), intro_p._p)
                if table is not None: parent.insert(parent.index(placeholder_p._p), table._tbl)
                parent.insert(parent.index(placeholder_p._p), spacer_p._p)
            
            # Remove the placeholder finally
            parent.remove(placeholder_p._p)

        # 2. Pagamentos Avulsos (DAMs) Table
        dams_placeholder = next((p for p in doc.paragraphs if '###DAMS_TABLE_PLACEHOLDER###' in p.text), None)
        
        if dams_placeholder:
            dams_data = context.get('pagamentos_avulsos', [])
            
            if dams_data:
                parent = dams_placeholder._p.getparent()
                
                intro_text = "Além dos fatos supracitados, foram identificados os seguintes pagamentos via DAM (Documento de Arrecadação Municipal) baixados para o contribuinte:"
                intro_p = doc.add_paragraph(intro_text)
                
                table = create_dams_table(doc, dams_data)
                spacer_p = doc.add_paragraph()

                curr_idx = parent.index(dams_placeholder._p)
                parent.insert(curr_idx, intro_p._p)
                if table is not None:
                    parent.insert(curr_idx + 1, table._tbl)
                parent.insert(curr_idx + 2, spacer_p._p)
                
                parent.remove(dams_placeholder._p)
                logging.info("Report Gen: Inserted DAMs table.")
            else:
                _delete_paragraph(dams_placeholder)
                logging.info("Report Gen: Removed DAMs placeholder (No data).")

        # 3. Conclusion Table
        conclusion_p = next((p for p in doc.paragraphs if '###CONCLUSION_TABLE_PLACEHOLDER###' in p.text), None)
        if conclusion_p:
            parent = conclusion_p._p.getparent()
            summary_table = create_conclusion_table(doc, context.get('summary'), idd_mode=idd_mode)
            if summary_table is not None: parent.insert(parent.index(conclusion_p._p), summary_table._tbl)
            parent.remove(conclusion_p._p)

        # ✅ Save DOCX
        _safe_save(doc, output_path)
        
        # ✅ NEW: Auto-Convert to PDF
        convert_to_pdf(output_path)

    except Exception as e:
        logging.exception(f"Erro ao gerar relatório: {e}")
        raise

def generate_simple_document(context, template_path, output_path):
    try:
        if not os.path.exists(template_path): raise FileNotFoundError(f"Template not found: {template_path}")
        
        doc_tpl = DocxTemplate(template_path)
        jinja_env = Environment(undefined=ChainableUndefined)
        doc_tpl.tpl_jinja_env = jinja_env
        
        custom_general_texts = get_custom_general_texts()
        rendered_general_texts = {}
        
        for key, template_string in custom_general_texts.items():
            try:
                template = jinja_env.from_string(template_string)
                rendered_general_texts[key] = template.render(context)
            except Exception:
                rendered_general_texts[key] = f"[[ERRO: {key}]]"
        
        render_context = context.copy()
        render_context.update(rendered_general_texts) 

        raw_idd_mode = context.get('idd_mode', False)
        idd_mode = str(raw_idd_mode).lower() in ('true', '1', 't') if isinstance(raw_idd_mode, str) else bool(raw_idd_mode)
        
        if idd_mode:
            if 'TITULO_DOCUMENTO_IDD' in rendered_general_texts:
                render_context['titulo_documento'] = rendered_general_texts['TITULO_DOCUMENTO_IDD']
            elif 'TITULO_DOCUMENTO_IDD' in DEFAULT_GENERAL_TEXTS:
                render_context['titulo_documento'] = DEFAULT_GENERAL_TEXTS['TITULO_DOCUMENTO_IDD']
            else:
                render_context['titulo_documento'] = "Informação Fiscal - Operação IDD"

            if 'I_INTRO_IDD' in rendered_general_texts:
                render_context['I_INTRO'] = rendered_general_texts['I_INTRO_IDD']
            elif 'I_INTRO_IDD' in DEFAULT_GENERAL_TEXTS:
                try:
                    tmpl = jinja_env.from_string(DEFAULT_GENERAL_TEXTS['I_INTRO_IDD'])
                    render_context['I_INTRO'] = tmpl.render(context)
                except:
                    render_context['I_INTRO'] = DEFAULT_GENERAL_TEXTS['I_INTRO_IDD']
        else:
             if 'TITULO_DOCUMENTO_AUTO' in rendered_general_texts:
                render_context['titulo_documento'] = rendered_general_texts['TITULO_DOCUMENTO_AUTO']
             elif 'TITULO_DOCUMENTO_AUTO' in DEFAULT_GENERAL_TEXTS:
                render_context['titulo_documento'] = DEFAULT_GENERAL_TEXTS['TITULO_DOCUMENTO_AUTO']

        render_context.pop('autos', None)
        render_context.pop('summary', None)
        
        doc_tpl.render(render_context)
        
        # ✅ Save DOCX
        _safe_save(doc_tpl, output_path)
        
        # ✅ NEW: Auto-Convert to PDF
        convert_to_pdf(output_path)

    except Exception as be:
        logging.error(f"Simple Doc CRASH: {be}")
        raise be