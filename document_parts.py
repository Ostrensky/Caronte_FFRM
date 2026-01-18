# --- FILE: document_parts.py ---

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt # ✅ NEW: Required for font sizing
import pandas as pd
import logging
import numpy as np
from app.config import get_custom_auto_texts

# ... (Previous helper functions remain unchanged: _format_currency_brl, format_invoice_numbers, etc.) ...

def _format_currency_brl(value):
    """Helper function to format a number into BRL currency string (1.234,56)."""
    if isinstance(value, (int, float)):
        return f"R$ {value:,.2f}".replace(",", "v").replace(".", ",").replace("v", ".")
    return str(value)

def _format_currency_brl_plain(value):
    """Helper function to format a number into BRL number string (1.234,56)."""
    if isinstance(value, (int, float)):
        return f"{value:,.2f}".replace(",", "v").replace(".", ",").replace("v", ".")
    return str(value)

def format_invoice_numbers(invoice_numbers):
    if isinstance(invoice_numbers, np.ndarray):
        if invoice_numbers.size == 0:
            return ""
    elif not invoice_numbers:
         return ""

    numeric_invoices = []
    for n in invoice_numbers:
        try:
            val = int(n)
            numeric_invoices.append(val)
        except (ValueError, TypeError):
            continue
            
    numeric_invoices = sorted(list(set(numeric_invoices)))

    if not numeric_invoices:
        return ", ".join(sorted(list(set([str(x) for x in invoice_numbers]))))

    ranges = []
    start_range = numeric_invoices[0]

    for i in range(1, len(numeric_invoices)):
        if numeric_invoices[i] != numeric_invoices[i-1] + 1:
            end_range = numeric_invoices[i-1]
            if start_range == end_range:
                ranges.append(str(start_range))
            else:
                ranges.append(f"{start_range} a {end_range}")
            start_range = numeric_invoices[i]

    end_range = numeric_invoices[-1]
    if start_range == end_range:
        ranges.append(str(start_range))
    else:
        ranges.append(f"{start_range} a {end_range}")

    return ", ".join(ranges)

def _safe_period_string(df):
    if 'DATA EMISSÃO' not in df.columns or df.empty or df['DATA EMISSÃO'].dropna().empty:
        return "[Período Indisponível]"
    min_date = df['DATA EMISSÃO'].min()
    max_date = df['DATA EMISSÃO'].max()
    if pd.isna(min_date) or pd.isna(max_date):
        return "[Período Indisponível]"
    min_str = min_date.strftime("%m/%Y")
    max_str = max_date.strftime("%m/%Y")
    return min_str if min_str == max_str else f"{min_str} a {max_str}"

def formatar_motivo_detalhado(motivo_data, custom_auto_texts=None):
    if custom_auto_texts is None:
        custom_auto_texts = get_custom_auto_texts()

    tipo = motivo_data.get('tipo', 'DEFAULT_AUTO_FALLBACK')
    fallback_template = custom_auto_texts.get("DEFAULT_AUTO_FALLBACK", "{{ texto_simples }}")
    template_text = custom_auto_texts.get(tipo, fallback_template)

    replacements = {
        "{{ correct_aliquota_str }}": str(motivo_data.get('aliquota_correta', '[N/A]')),
        "{{ texto_simples }}": motivo_data.get('texto_simples', 'Motivo não especificado.')
    }

    formatted_text = template_text
    for placeholder, value in replacements.items():
        if placeholder in formatted_text:
             formatted_text = formatted_text.replace(placeholder, value)

    return formatted_text

def formatar_texto_multa(auto_data, df_invoices, valor_multa_str, multas_list=None):
    try:
        if df_invoices is None or df_invoices.empty:
            return "Nenhuma nota fiscal associada para gerar o texto da multa."

        df_safe = df_invoices.copy()
        if 'NÚMERO' not in df_safe.columns:
             return "Coluna 'NÚMERO' ausente nos dados das notas fiscais."
        df_safe['NÚMERO'] = df_safe['NÚMERO'].astype(str).str.strip()

        rule_to_motive_text = {
            'IDD (Não Pago)': 'à apuração do imposto devido',
            'Dedução indevida': 'a dedução de valores da base de cálculo',
            'Regime incorreto': 'ao regime tributário',
            'Isenção/Imunidade Indevida': 'a isenção/imunidade quando permitida pela legislação',
            'Natureza da Operação Incompatível': 'ao local da incidência do imposto',
            'Local da incidência incorreto': 'ao local da incidência do imposto',
            'Retenção na Fonte (Verificar)': 'a retenção na fonte quando permitida pela legislação',
            'Alíquota Incorreta': 'a alíquota'
        }
        alinea_map = {
            'IDD (Não Pago)': 'a', 'Dedução indevida': 'b', 'Regime incorreto': 'd',
            'Isenção/Imunidade Indevida': 'e', 'Natureza da Operação Incompatível': 'h',
            'Local da incidência incorreto': 'h', 'Retenção na Fonte (Verificar)': 'i',
            'Alíquota Incorreta': 'c'
        }

        aggregated_infractions = {}
        instrumental_causes = list(rule_to_motive_text.keys())
        instrumental_causes.remove('IDD (Não Pago)') 

        for _, row in df_safe.iterrows():
            if not isinstance(row.get('broken_rule_details'), list):
                continue
            
            instrumental_details_for_row = [
                detail for detail in set(row['broken_rule_details'])
                if any(detail.startswith(cause) for cause in instrumental_causes)
            ]

            for detail in instrumental_details_for_row:
                for cause in instrumental_causes:
                    if detail.startswith(cause):
                        if cause not in aggregated_infractions:
                            aggregated_infractions[cause] = set()
                        aggregated_infractions[cause].add(row['NÚMERO'])
                        break

        if not aggregated_infractions:
            return "Não foi constatada a emissão incorreta de notas no período em análise. Não há multa por descumprimento de dever instrumental."

        main_text_parts = []
        all_motives = list(aggregated_infractions.keys())

        for cause, invoice_numbers_set in aggregated_infractions.items():
            invoice_numbers = list(invoice_numbers_set)
            group_df = df_safe[df_safe['NÚMERO'].isin(invoice_numbers)]
            nfs_e_numeros = format_invoice_numbers(invoice_numbers)
            periodo = _safe_period_string(group_df)
            motivo_texto_pt = rule_to_motive_text.get(cause, "a dados incorretos")
            
            line = f"- NFS-e de nº (s) {nfs_e_numeros} com dados incorretos referentes {motivo_texto_pt} no período de {periodo}."
            if cause.startswith('Retenção na Fonte'):
                 line += " Em desacordo com o disposto no artigo 8º e 8º-A da Lei Complementar nº 40/2001 e alterações."
            main_text_parts.append(line)
        
        if len(main_text_parts) == 1:
            single_cause_text = main_text_parts[0].replace("- ", "").replace(".", "")
            body_text = (f"O sujeito passivo acima identificado emitiu {single_cause_text}, "
                         "caracterizando-se como descumprimento de dever instrumental.")
        else:
            intro = ("O sujeito passivo acima identificado emitiu Notas Fiscais Eletrônicas (NFS-e) de "
                     "prestação de serviços com dados incorretos, caracterizando-se como "
                     "descumprimento de dever instrumental. São elas:")
            body_text = intro + "\n" + "\n".join(main_text_parts)

        alineas_encontradas = set()
        for motive in all_motives:
            for key, alinea in alinea_map.items():
                if motive.startswith(key):
                    alineas_encontradas.add(alinea)
            if motive.startswith('Isenção/Imunidade Indevida'):
                alineas_encontradas.add('f')

        alinea_texto = ""
        if alineas_encontradas:
            sorted_alineas = sorted(list(alineas_encontradas))
            if len(sorted_alineas) == 1:
                alinea_texto = f"alínea “{sorted_alineas[0]}”,"
            else:
                alineas_formatadas = [f"“{a}”" for a in sorted_alineas]
                alinea_texto = f"alíneas {', '.join(alineas_formatadas[:-1])} e {alineas_formatadas[-1]},"
        
        is_retencao = any(m.startswith('Retenção na Fonte') for m in all_motives)
        texto_final_parte1 = (
            f"\n\nConstituindo a ocorrência infração ao disposto no artigo 12, § 1º, inciso III, {alinea_texto} da Lei Complementar nº 73/2009 e alterações"
            f"{', e Decreto Municipal 1.712/2020, art. 10, caput e §1º' if is_retencao else ''}."
        )
        texto_final_parte2 = (
            f"\n\nSendo lavrado o presente auto de infração, em 2 (duas) vias, com a cominação da multa correspondente a importância de {valor_multa_str}."
            "\n\nMulta calculada conforme proporção de ocorrências definida no § 2º do artigo 12 da Lei Complementar nº 73/2009 e alterações."
            "\n\nMontante atualizado nos termos do Decreto nº 1.928 de 10 de dezembro de 2024."
            "\n\nO valor da multa acima descrita será reduzido em 50 por cento para a ME/EPP optante pelo Simples Nacional, "
            "conforme o artigo 38-B, inciso II, da Lei Complementar nº123/2006, caso o pagamento seja realizado "
            "dentro de 30 (trinta) dias contados da data da ciência."
        )

        final_text = body_text + texto_final_parte1 + texto_final_parte2

        # ✅ NEW: Append multiple fines explanation if applicable
        if multas_list and len(multas_list) > 1:
            # Extract non-empty numbers
            fine_numbers = [str(m.get('number', '')).strip() for m in multas_list if str(m.get('number', '')).strip()]
            
            if fine_numbers:
                if len(fine_numbers) == 1:
                    ids_str = fine_numbers[0]
                else:
                    ids_str = ", ".join(fine_numbers[:-1]) + " e " + fine_numbers[-1]
                
                final_text += f"\n\nComo as ocorrências descritas acima compreenderam mais de um exercício fiscal, foram emitidas as multas {ids_str}."

        return final_text
    
    except Exception as e:
        logging.exception("Error in formatar_texto_multa")
        return f"ERRO INESPERADO AO GERAR TEXTO DA MULTA: {e}"

# ✅ --- START: New Helper to force font size ---
def _set_table_font_size(table, size_pt):
    """Iterates through all cells in a table and sets the font size."""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(size_pt)
# ✅ --- END: New Helper ---

def create_table_for_auto(doc, auto_data, idd_mode=False):
    table = None
    has_das = auto_data.get('tem_pagamento_das', False)
    has_dam = auto_data.get('tem_pagamento_dam', False)

    dados_anuais_para_tabela = auto_data.get('dados_anuais', [])
    
    if not dados_anuais_para_tabela:
        return None

    total_pago_idd = sum(float(mes.get('iss_declarado_pago', 0.0)) for mes in dados_anuais_para_tabela)
    has_pagamento_idd = total_pago_idd > 0.01

    if has_das and has_dam:
        table = _create_table_das_dam(doc, auto_data, dados_anuais_para_tabela, has_pagamento_idd, idd_mode)
    elif has_das:
        table = _create_table_das_only(doc, auto_data, dados_anuais_para_tabela, has_pagamento_idd, idd_mode)
    elif has_dam:
        table = _create_table_dam_only(doc, auto_data, dados_anuais_para_tabela, has_pagamento_idd, idd_mode)
    else:
        # Simplificado usually stays at default font size (10 or 11)
        table = _create_table_simplificado(doc, auto_data, dados_anuais_para_tabela, has_pagamento_idd, idd_mode)

    if table:
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return table

def _create_table_das_dam(doc, auto_data, dados_anuais_filtrados, has_pagamento_idd, idd_mode=False):
    cols = 18 if has_pagamento_idd else 15
    table = doc.add_table(rows=2, cols=cols, style='Table Grid')
    hdr1, hdr2 = table.rows
    
    hdr1.cells[0].merge(hdr2.cells[0]).text = 'MÊS/ANO'
    
    hdr1.cells[1].merge(hdr1.cells[3]).text = 'APURAÇÃO DO ISS'
    hdr2.cells[1].text = 'Base de Cálculo'
    hdr2.cells[2].text = 'Alíquota'
    hdr2.cells[3].text = 'ISS Apurado'
    
    col_offset = 0
    if has_pagamento_idd:
        hdr1.cells[4].merge(hdr1.cells[6]).text = 'PAGAMENTOS (IDD)'
        hdr2.cells[4].text = 'Base de Cálculo'
        hdr2.cells[5].text = 'Alíquota'
        hdr2.cells[6].text = 'ISS Pago'
        col_offset = 3 
    
    hdr1.cells[4+col_offset].merge(hdr1.cells[6+col_offset]).text = 'ISS PAGO POR DAS'
    hdr2.cells[4+col_offset].text = 'Base de Cálculo'
    hdr2.cells[5+col_offset].text = 'Alíquota'
    hdr2.cells[6+col_offset].text = 'ISS Pago'

    hdr1.cells[7+col_offset].merge(hdr1.cells[9+col_offset]).text = 'ISS PAGO POR DAM'
    hdr2.cells[7+col_offset].text = 'Base de Cálculo'
    hdr2.cells[8+col_offset].text = 'Alíquota'
    hdr2.cells[9+col_offset].text = 'ISS Pago'

    titulo_coluna = f"IDD Nº {auto_data['numero']}" if idd_mode else f"AUTO DE INFRAÇÃO Nº {auto_data['numero']}"

    hdr1.cells[10+col_offset].merge(hdr1.cells[12+col_offset]).text = titulo_coluna
    hdr2.cells[10+col_offset].text = 'Base de Cálculo'
    hdr2.cells[11+col_offset].text = 'Alíquota'
    hdr2.cells[12+col_offset].text = 'ISS constituído'
    
    hdr1.cells[13+col_offset].merge(hdr1.cells[14+col_offset]).text = 'IDENTIFICAÇÃO'
    hdr2.cells[13+col_offset].text = 'DAS'
    hdr2.cells[14+col_offset].text = 'DAM'
    
    for mes in dados_anuais_filtrados:
        row = table.add_row().cells
        row[0].text = mes['mes_ano']
        
        row[1].text = _format_currency_brl_plain(mes['base_calculo'])
        row[2].text = str(mes['aliquota_op']) 
        row[3].text = _format_currency_brl_plain(mes['iss_apurado_bruto']) 
        
        if has_pagamento_idd:
            row[4].text = _format_currency_brl_plain(mes['base_calculo'])
            row[5].text = str(mes['aliquota_declarada']) 
            row[6].text = _format_currency_brl_plain(mes['iss_declarado_pago']) 
        
        row[4+col_offset].text = _format_currency_brl_plain(mes['base_calculo'])
        row[5+col_offset].text = str(mes['das_aliquota']) 
        row[6+col_offset].text = _format_currency_brl_plain(mes['das_iss_pago']) 
        
        row[7+col_offset].text = _format_currency_brl_plain(mes['base_calculo'])
        row[8+col_offset].text = str(mes['dam_aliquota'])
        row[9+col_offset].text = _format_currency_brl_plain(mes['dam_iss_pago'])
        
        row[10+col_offset].text = _format_currency_brl_plain(mes['base_calculo_op'])
        if float(mes.get('base_calculo_op', 0.0)) > 0.001:
            aliquota_final_efetiva = (float(mes['iss_apurado_op']) / float(mes['base_calculo_op'])) * 100.0
            row[11+col_offset].text = f"{aliquota_final_efetiva:.2f}%"
        else:
            row[11+col_offset].text = "-"
        row[12+col_offset].text = _format_currency_brl_plain(mes['iss_apurado_op'])
        
        row[13+col_offset].text = str(mes['das_identificacao'])
        row[14+col_offset].text = str(mes['dam_identificacao'])

    total = table.add_row().cells; total[0].text = 'TOTAL'
    totais = auto_data['totais']
    
    total[1].text = _format_currency_brl_plain(totais['base_calculo'])
    total[2].text = "-"
    total[3].text = _format_currency_brl_plain(totais['iss_apurado_bruto'])
    if has_pagamento_idd:
        total[4].text = _format_currency_brl_plain(totais['base_calculo'])
        total[5].text = "-"
        total[6].text = _format_currency_brl_plain(totais['iss_declarado_pago'])
    
    total[4+col_offset].text = _format_currency_brl_plain(totais['base_calculo'])
    total[5+col_offset].text = "-"
    total[6+col_offset].text = _format_currency_brl_plain(totais['das_iss_pago'])
    
    total[7+col_offset].text = _format_currency_brl_plain(totais['base_calculo'])
    total[8+col_offset].text = "-"
    total[9+col_offset].text = _format_currency_brl_plain(totais['dam_iss_pago'])
    
    total[10+col_offset].text = _format_currency_brl_plain(totais['base_calculo_op'])
    total[11+col_offset].text = "-"
    total[12+col_offset].text = _format_currency_brl_plain(totais['iss_apurado_op'])
    
    total[13+col_offset].text = "-"
    total[14+col_offset].text = "-"

    # ✅ Apply 8pt font size
    _set_table_font_size(table, 8)
    return table

def _create_table_das_only(doc, auto_data, dados_anuais_filtrados, has_pagamento_idd, idd_mode=False):
    cols = 15 if has_pagamento_idd else 12
    table = doc.add_table(rows=2, cols=cols, style='Table Grid')
    hdr1, hdr2 = table.rows

    hdr1.cells[0].merge(hdr2.cells[0]).text = 'MÊS/ANO'
    
    hdr1.cells[1].merge(hdr1.cells[3]).text = 'APURAÇÃO DO ISS'
    hdr2.cells[1].text = 'Base de Cálculo'
    hdr2.cells[2].text = 'Alíquota'
    hdr2.cells[3].text = 'ISS Apurado'
    
    col_offset = 0
    if has_pagamento_idd:
        hdr1.cells[4].merge(hdr1.cells[6]).text = 'PAGAMENTOS (IDD)'
        hdr2.cells[4].text = 'Base de Cálculo'
        hdr2.cells[5].text = 'Alíquota'
        hdr2.cells[6].text = 'ISS Pago'
        col_offset = 3
    
    hdr1.cells[4+col_offset].merge(hdr1.cells[6+col_offset]).text = 'ISS PAGO POR DAS'
    hdr2.cells[4+col_offset].text = 'Base de Cálculo'
    hdr2.cells[5+col_offset].text = 'Alíquota'
    hdr2.cells[6+col_offset].text = 'ISS Pago'

    titulo_coluna = f"IDD Nº {auto_data['numero']}" if idd_mode else f"AUTO DE INFRAÇÃO Nº {auto_data['numero']}"

    hdr1.cells[7+col_offset].merge(hdr1.cells[9+col_offset]).text = titulo_coluna
    hdr2.cells[7+col_offset].text = 'Base de Cálculo'
    hdr2.cells[8+col_offset].text = 'Alíquota'
    hdr2.cells[9+col_offset].text = 'ISS constituído'
    
    hdr1.cells[10+col_offset].merge(hdr2.cells[10+col_offset]).text = 'IDENTIFICAÇÃO'
    hdr1.cells[11+col_offset].merge(hdr2.cells[11+col_offset]).text = 'DAS'
    hdr1.cells[10+col_offset].merge(hdr1.cells[11+col_offset])

    for mes in dados_anuais_filtrados:
        row = table.add_row().cells
        row[0].text = mes['mes_ano']
        
        row[1].text = _format_currency_brl_plain(mes['base_calculo'])
        row[2].text = str(mes['aliquota_op']) 
        row[3].text = _format_currency_brl_plain(mes['iss_apurado_bruto']) 
        
        if has_pagamento_idd:
            row[4].text = _format_currency_brl_plain(mes['base_calculo'])
            row[5].text = str(mes['aliquota_declarada']) 
            row[6].text = _format_currency_brl_plain(mes['iss_declarado_pago']) 
        
        row[4+col_offset].text = _format_currency_brl_plain(mes['base_calculo'])
        row[5+col_offset].text = str(mes['das_aliquota']) 
        row[6+col_offset].text = _format_currency_brl_plain(mes['das_iss_pago']) 
        
        row[7+col_offset].text = _format_currency_brl_plain(mes['base_calculo_op'])
        if float(mes.get('base_calculo_op', 0.0)) > 0.001:
            aliquota_final_efetiva = (float(mes['iss_apurado_op']) / float(mes['base_calculo_op'])) * 100.0
            row[8+col_offset].text = f"{aliquota_final_efetiva:.2f}%"
        else:
            row[8+col_offset].text = "-"
        row[9+col_offset].text = _format_currency_brl_plain(mes['iss_apurado_op'])
        
        row[10+col_offset].text = str(mes['das_identificacao'])
        row[10+col_offset].merge(row[11+col_offset]) 

    total = table.add_row().cells
    total[0].text = 'TOTAL'
    totais = auto_data['totais']
    
    total[1].text = _format_currency_brl_plain(totais['base_calculo'])
    total[2].text = "-"
    total[3].text = _format_currency_brl_plain(totais['iss_apurado_bruto'])
    if has_pagamento_idd:
        total[4].text = _format_currency_brl_plain(totais['base_calculo'])
        total[5].text = "-"
        total[6].text = _format_currency_brl_plain(totais['iss_declarado_pago'])
    
    total[4+col_offset].text = _format_currency_brl_plain(totais['base_calculo'])
    total[5+col_offset].text = "-"
    total[6+col_offset].text = _format_currency_brl_plain(totais['das_iss_pago'])
    
    total[7+col_offset].text = _format_currency_brl_plain(totais['base_calculo_op'])
    total[8+col_offset].text = "-"
    total[9+col_offset].text = _format_currency_brl_plain(totais['iss_apurado_op'])
    
    total[10+col_offset].text = "-"
    total[10+col_offset].merge(total[11+col_offset])

    # ✅ Apply 8pt font size
    _set_table_font_size(table, 8)
    return table

def _create_table_dam_only(doc, auto_data, dados_anuais_filtrados, has_pagamento_idd, idd_mode=False):
    cols = 15 if has_pagamento_idd else 12
    table = doc.add_table(rows=2, cols=cols, style='Table Grid')
    hdr1, hdr2 = table.rows
    
    hdr1.cells[0].merge(hdr2.cells[0]).text = 'MÊS/ANO'
    
    hdr1.cells[1].merge(hdr1.cells[3]).text = 'APURAÇÃO DO ISS'
    hdr2.cells[1].text = 'Base de Cálculo'
    hdr2.cells[2].text = 'Alíquota'
    hdr2.cells[3].text = 'ISS Apurado'
    
    col_offset = 0
    if has_pagamento_idd:
        hdr1.cells[4].merge(hdr1.cells[6]).text = 'PAGAMENTOS (IDD)'
        hdr2.cells[4].text = 'Base de Cálculo'
        hdr2.cells[5].text = 'Alíquota'
        hdr2.cells[6].text = 'ISS Pago'
        col_offset = 3
    
    hdr1.cells[4+col_offset].merge(hdr1.cells[6+col_offset]).text = 'ISS PAGO POR DAM'
    hdr2.cells[4+col_offset].text = 'Base de Cálculo'
    hdr2.cells[5+col_offset].text = 'Alíquota'
    hdr2.cells[6+col_offset].text = 'ISS Pago'

    titulo_coluna = f"IDD Nº {auto_data['numero']}" if idd_mode else f"AUTO DE INFRAÇÃO Nº {auto_data['numero']}"
    
    hdr1.cells[7+col_offset].merge(hdr1.cells[9+col_offset]).text = titulo_coluna
    hdr2.cells[7+col_offset].text = 'Base de Cálculo'
    hdr2.cells[8+col_offset].text = 'Alíquota'
    hdr2.cells[9+col_offset].text = 'ISS constituído'
    
    hdr1.cells[10+col_offset].merge(hdr2.cells[10+col_offset]).text = 'IDENTIFICAÇÃO'
    hdr1.cells[11+col_offset].merge(hdr2.cells[11+col_offset]).text = 'DAM'
    hdr1.cells[10+col_offset].merge(hdr1.cells[11+col_offset])

    for mes in dados_anuais_filtrados:
        row = table.add_row().cells
        row[0].text = mes['mes_ano']
        
        row[1].text = _format_currency_brl_plain(mes['base_calculo'])
        row[2].text = str(mes['aliquota_op']) 
        row[3].text = _format_currency_brl_plain(mes['iss_apurado_bruto']) 
        
        if has_pagamento_idd:
            row[4].text = _format_currency_brl_plain(mes['base_calculo'])
            row[5].text = str(mes['aliquota_declarada']) 
            row[6].text = _format_currency_brl_plain(mes['iss_declarado_pago']) 
        
        row[4+col_offset].text = _format_currency_brl_plain(mes['base_calculo'])
        row[5+col_offset].text = str(mes['dam_aliquota']) 
        row[6+col_offset].text = _format_currency_brl_plain(mes['dam_iss_pago'])
        
        row[7+col_offset].text = _format_currency_brl_plain(mes['base_calculo_op'])
        if float(mes.get('base_calculo_op', 0.0)) > 0.001:
            aliquota_final_efetiva = (float(mes['iss_apurado_op']) / float(mes['base_calculo_op'])) * 100.0
            row[8+col_offset].text = f"{aliquota_final_efetiva:.2f}%"
        else:
            row[8+col_offset].text = "-"
        row[9+col_offset].text = _format_currency_brl_plain(mes['iss_apurado_op'])
        
        row[10+col_offset].text = str(mes['dam_identificacao'])
        row[10+col_offset].merge(row[11+col_offset])

    total = table.add_row().cells
    total[0].text = 'TOTAL'
    totais = auto_data['totais']
    
    total[1].text = _format_currency_brl_plain(totais['base_calculo'])
    total[2].text = "-"
    total[3].text = _format_currency_brl_plain(totais['iss_apurado_bruto'])
    if has_pagamento_idd:
        total[4].text = _format_currency_brl_plain(totais['base_calculo'])
        total[5].text = "-"
        total[6].text = _format_currency_brl_plain(totais['iss_declarado_pago'])
    
    total[4+col_offset].text = _format_currency_brl_plain(totais['base_calculo'])
    total[5+col_offset].text = "-"
    total[6+col_offset].text = _format_currency_brl_plain(totais['dam_iss_pago'])
    
    total[7+col_offset].text = _format_currency_brl_plain(totais['base_calculo_op'])
    total[8+col_offset].text = "-"
    total[9+col_offset].text = _format_currency_brl_plain(totais['iss_apurado_op'])
    
    total[10+col_offset].text = "-"
    total[10+col_offset].merge(total[11+col_offset])

    # ✅ Apply 8pt font size
    _set_table_font_size(table, 8)
    return table

def _create_table_simplificado(doc, auto_data, dados_anuais_filtrados, has_pagamento_idd, idd_mode=False):
    # This function remains without specific font sizing (uses default)
    # as the prompt specifically asked to change it "if it has dams/das".
    cols = 10 if has_pagamento_idd else 7
    table = doc.add_table(rows=2, cols=cols, style='Table Grid')
    hdr1, hdr2 = table.rows
    
    hdr1.cells[0].merge(hdr2.cells[0]).text = 'MÊS/ANO'
    
    hdr1.cells[1].merge(hdr1.cells[3]).text = 'APURAÇÃO DO ISS'
    hdr2.cells[1].text = 'Base de Cálculo'
    hdr2.cells[2].text = 'Alíquota'
    hdr2.cells[3].text = 'ISS Apurado'
    
    col_offset = 0
    if has_pagamento_idd:
        hdr1.cells[4].merge(hdr1.cells[6]).text = 'PAGAMENTOS (IDD)'
        hdr2.cells[4].text = 'Base de Cálculo'
        hdr2.cells[5].text = 'Alíquota'
        hdr2.cells[6].text = 'ISS Pago'
        col_offset = 3

    titulo_coluna = f"IDD Nº {auto_data['numero']}" if idd_mode else f"AUTO DE INFRAÇÃO Nº {auto_data['numero']}"

    hdr1.cells[4+col_offset].merge(hdr1.cells[6+col_offset]).text = titulo_coluna
    hdr2.cells[4+col_offset].text = 'Base de Cálculo'
    hdr2.cells[5+col_offset].text = 'Alíquota'
    hdr2.cells[6+col_offset].text = 'ISS constituído'
    
    for mes in dados_anuais_filtrados:
        row = table.add_row().cells
        row[0].text = mes['mes_ano']
        
        row[1].text = _format_currency_brl_plain(mes['base_calculo'])
        row[2].text = str(mes['aliquota_op']) 
        row[3].text = _format_currency_brl_plain(mes['iss_apurado_bruto']) 
        
        if has_pagamento_idd:
            row[4].text = _format_currency_brl_plain(mes['base_calculo'])
            row[5].text = str(mes['aliquota_declarada']) 
            row[6].text = _format_currency_brl_plain(mes['iss_declarado_pago']) 
        
        row[4+col_offset].text = _format_currency_brl_plain(mes['base_calculo_op'])
        if float(mes.get('base_calculo_op', 0.0)) > 0.001:
            aliquota_final_efetiva = (float(mes['iss_apurado_op']) / float(mes['base_calculo_op'])) * 100.0
            row[5+col_offset].text = f"{aliquota_final_efetiva:.2f}%"
        else:
            row[5+col_offset].text = "-"
        row[6+col_offset].text = _format_currency_brl_plain(mes['iss_apurado_op']) 
    
    total = table.add_row().cells
    total[0].text = 'TOTAL'
    totais = auto_data['totais']
    
    total[1].text = _format_currency_brl_plain(totais['base_calculo'])
    total[2].text = "-"
    total[3].text = _format_currency_brl_plain(totais['iss_apurado_bruto'])
    if has_pagamento_idd:
        total[4].text = _format_currency_brl_plain(totais['base_calculo'])
        total[5].text = "-"
        total[6].text = _format_currency_brl_plain(totais['iss_declarado_pago'])
    
    total[4+col_offset].text = _format_currency_brl_plain(totais['base_calculo_op'])
    total[5+col_offset].text = "-"
    total[6+col_offset].text = _format_currency_brl_plain(totais['iss_apurado_op'])
    
    return table